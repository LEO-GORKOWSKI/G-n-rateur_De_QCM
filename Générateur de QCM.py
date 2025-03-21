import random
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Cm

class GenerateurQCM:
    """
    Classe principale pour générer des QCM à partir d'un fichier de questions.
    """
    def __init__(self, fichier_questions, nb_eleves, nb_questions, graine):
        """
        Initialise le générateur de QCM.
        
        Args:
            fichier_questions (str): Chemin vers le fichier contenant les questions
            nb_eleves (int): Nombre d'élèves (donc de sujets à créer)
            nb_questions (int): Nombre de questions par sujet (5, 10, 15 ou 20)
            graine (int): Graine pour la génération aléatoire
        """
        self.fichier_questions = fichier_questions
        self.nb_eleves = nb_eleves
        self.nb_questions = nb_questions
        self.graine = graine
        
        # Fixer la graine aléatoire
        random.seed(graine)
        
        # Charger les questions
        self.questions = self.charger_questions()
        
        # Vérifier qu'il y a assez de questions
        if len(self.questions) < nb_questions:
            raise ValueError(f"Le fichier ne contient que {len(self.questions)} questions, mais {nb_questions} sont demandées.")
        
        # Pour suivre quelles questions apparaissent dans quels sujets
        self.questions_par_sujet = {}
        self.sujets_par_question = {i: [] for i in range(len(self.questions))}
        
    def charger_questions(self):
        """
        Charge les questions depuis le fichier.
        
        Returns:
            list: Liste de dictionnaires contenant les questions, réponses et solution
        """
        questions = []
        
        with open(self.fichier_questions, 'r', encoding='utf-8') as f:
            lignes = f.readlines()
            
            i = 0
            while i < len(lignes):
                # Si on est à la fin du fichier ou devant une ligne vide, passer
                if i >= len(lignes) or not lignes[i].strip():
                    i += 1
                    continue
                
                # Récupérer la question
                question = lignes[i].strip()
                
                # Récupérer les 4 réponses
                reponses = []
                for j in range(1, 5):
                    if i + j < len(lignes) and lignes[i + j].strip():
                        reponses.append(lignes[i + j].strip())
                    else:
                        # Si pas assez de réponses, passer cette question
                        break
                
                # Récupérer la solution
                if i + 5 < len(lignes) and lignes[i + 5].strip():
                    try:
                        solution = int(lignes[i + 5].strip())
                        
                        # Vérifier que la solution est valide (entre 1 et 4)
                        if 1 <= solution <= 4:
                            # Ajouter la question
                            questions.append({
                                'question': question,
                                'reponses': reponses,
                                'solution': solution
                            })
                    except ValueError:
                        # Si la solution n'est pas un entier, passer cette question
                        pass
                
                # Passer à la question suivante (question + 4 réponses + solution + ligne vide)
                i += 7
        
        return questions
    
    def generer_sujet(self, num_sujet):
        """
        Génère un sujet pour un élève.
        
        Args:
            num_sujet (int): Numéro du sujet à générer
            
        Returns:
            tuple: (liste de questions pour ce sujet, 
                   liste des réponses correctes sous forme de lettres)
        """
        # Sélectionner aléatoirement les questions pour ce sujet
        indices_questions = random.sample(range(len(self.questions)), self.nb_questions)
        
        # Enregistrer quelles questions apparaissent dans ce sujet
        self.questions_par_sujet[num_sujet] = indices_questions
        
        # Mettre à jour les sujets dans lesquels chaque question apparaît
        for idx in indices_questions:
            self.sujets_par_question[idx].append(num_sujet)
        
        questions_sujet = []
        reponses_correctes = []
        
        for i, idx in enumerate(indices_questions):
            question = self.questions[idx]
            
            # Copier les réponses pour pouvoir les mélanger
            reponses = question['reponses'].copy()
            solution_originale = question['solution'] - 1  # Convertir 1-4 en 0-3
            
            # Mélanger les réponses
            random.shuffle(reponses)
            
            # Trouver la nouvelle position de la bonne réponse
            nouvelle_solution = reponses.index(question['reponses'][solution_originale])
            
            # Convertir l'indice (0-3) en lettre (a-d)
            lettre_solution = chr(97 + nouvelle_solution)  # 'a', 'b', 'c' ou 'd'
            
            # Ajouter la question et les réponses mélangées au sujet
            questions_sujet.append({
                'numero': i + 1,
                'question': question['question'],
                'reponses': reponses,
                'solution': nouvelle_solution + 1  # Reconvertir 0-3 en 1-4
            })
            
            # Ajouter la lettre de la bonne réponse à la liste
            reponses_correctes.append(lettre_solution)
        
        return questions_sujet, reponses_correctes
    
    def generer_sujets(self):
        """
        Génère tous les sujets pour tous les élèves.
        
        Returns:
            dict: Dictionnaire de sujets par numéro de sujet
        """
        sujets = {}
        
        for i in range(self.nb_eleves):
            sujets[i] = self.generer_sujet(i)
        
        return sujets
    
    def generer_fichier_txt(self, sujets, nom_fichier="qcm_sujets.txt"):
        """
        Génère un fichier texte contenant tous les sujets.
        
        Args:
            sujets (dict): Dictionnaire de sujets par numéro de sujet
            nom_fichier (str): Nom du fichier à générer
        """
        with open(nom_fichier, 'w', encoding='utf-8') as f:
            for num_sujet, (questions, _) in sujets.items():
                f.write(f"SUJET {num_sujet}\n")
                f.write("\nNom prénom : _______________________________\n\n")
                
                # Consignes
                f.write("CONSIGNES :\n")
                f.write("- Entourez la lettre correspondant à la bonne réponse sur le tableau de réponses ci-dessous.\n")
                f.write("- Une seule réponse est correcte pour chaque question.\n")
                f.write("- Toute rature ou correction sur le tableau sera considérée comme une erreur.\n\n")
                
                # Tableau de réponses
                f.write("Tableau de réponses - Sujet " + str(num_sujet) + "\n")
                for i in range(1, self.nb_questions + 1):
                    if i == self.nb_questions // 2 + 1:
                        f.write("\n")  # Saut de ligne après la moitié des questions
                    f.write(f"{i} ")
                f.write("\n\n")
                
                # Questions
                for q in questions:
                    f.write(f"Question {q['numero']} : {q['question']}\n")
                    for j, reponse in enumerate(q['reponses']):
                        lettre = chr(97 + j)  # 'a', 'b', 'c' ou 'd'
                        f.write(f"{lettre}) {reponse}\n")
                    f.write("\n")
                
                f.write("\n---------------------------------------------------\n\n")
    
    def generer_fichier_correction(self, sujets, nom_fichier="qcm_corrections.txt"):
        """
        Génère un fichier de correction contenant les réponses correctes pour chaque sujet.
        
        Args:
            sujets (dict): Dictionnaire de sujets par numéro de sujet
            nom_fichier (str): Nom du fichier à générer
        """
        with open(nom_fichier, 'w', encoding='utf-8') as f:
            # Écrire les corrections pour chaque sujet
            for num_sujet, (_, reponses) in sujets.items():
                f.write(f"Sujet {num_sujet}\n")
                
                # Écrire les réponses en groupes de 5
                for i in range(0, len(reponses), 5):
                    groupe = reponses[i:i+5]
                    f.write(''.join(groupe) + ' ')
                f.write("\n\n")
            
            # Ajouter les informations supplémentaires pour le correcteur
            f.write("\n--- Informations supplémentaires pour le correcteur ---\n\n")
            
            # Quelles questions initiales ont été intégrées dans chaque sujet
            for num_sujet, indices in self.questions_par_sujet.items():
                indices_str = ','.join(str(i) for i in indices)
                f.write(f"Sujet {num_sujet} : questions {indices_str}\n")
            
            f.write("\n")
            
            # Dans quels sujets apparaît chaque question initiale
            for num_question, sujets_indices in self.sujets_par_question.items():
                if sujets_indices:  # Si la question apparaît dans au moins un sujet
                    sujets_str = ','.join(str(i) for i in sujets_indices)
                    f.write(f"Question {num_question} : sujets {sujets_str}\n")
    
    def generer_fichier_docx(self, sujets, nom_fichier="qcm_sujets.docx"):
        """
        Génère un fichier DOCX contenant tous les sujets.
        
        Args:
            sujets (dict): Dictionnaire de sujets par numéro de sujet
            nom_fichier (str): Nom du fichier à générer
        """
        document = Document()
        
        # Définir la mise en page : paysage, marges réduites
        section = document.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        
        # Réduire les marges
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        
        # Créer les sujets
        for num_sujet, (questions, _) in sujets.items():
            # Titre
            titre = document.add_paragraph()
            titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
            titre_run = titre.add_run(f"SUJET {num_sujet}")
            titre_run.bold = True
            titre_run.font.size = Pt(16)
            
            # Champ Nom/Prénom
            nom = document.add_paragraph()
            nom.add_run("Nom prénom : _______________________________")
            
            # Consignes
            consignes = document.add_paragraph()
            consignes.add_run("CONSIGNES :\n").bold = True
            consignes.add_run("- Entourez la lettre correspondant à la bonne réponse sur le tableau de réponses ci-dessous.\n")
            consignes.add_run("- Une seule réponse est correcte pour chaque question.\n")
            consignes.add_run("- Toute rature ou correction sur le tableau sera considérée comme une erreur.")
            
            # Tableau de réponses
            reponses_titre = document.add_paragraph()
            reponses_titre.add_run(f"Tableau de réponses - Sujet {num_sujet}").bold = True
            
            # Ajouter les numéros de question pour le tableau de réponses
            tableau_reponses = document.add_paragraph()
            for i in range(1, self.nb_questions + 1):
                if i == self.nb_questions // 2 + 1:
                    tableau_reponses.add_run("\n")  # Saut de ligne à la moitié
                tableau_reponses.add_run(f"{i} ")
            
            document.add_paragraph()  # Espace
            
            # Créer une section à deux colonnes pour les questions
            section = document.add_section(WD_SECTION.NEW_PAGE)
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width, section.page_height = section.page_height, section.page_width
            section.left_margin = Cm(1.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            
            # Configurer 2 colonnes
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.column_width = Inches(4)
            section.column_count = 2
            
            # Questions et réponses
            for q in questions:
                question_para = document.add_paragraph()
                question_para.add_run(f"Question {q['numero']} : {q['question']}").bold = True
                
                for j, reponse in enumerate(q['reponses']):
                    lettre = chr(97 + j)  # 'a', 'b', 'c' ou 'd'
                    reponse_para = document.add_paragraph()
                    reponse_para.add_run(f"{lettre}) {reponse}")
                
                document.add_paragraph()  # Espace entre les questions
            
            # Ajouter un saut de page après chaque sujet sauf le dernier
            if num_sujet < self.nb_eleves - 1:
                document.add_page_break()
        
        # Sauvegarder le document
        document.save(nom_fichier)
    
    